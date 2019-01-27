import fonctions
from selenium import webdriver
from selenium.webdriver.firefox.options import Options
import tkinter
from tkinter import filedialog
from docx import Document
import threading
import time


# region Cette partie permet de demander à l'utilisateur le fichier qu'il veux modifier
taille_min_mot = 6

file_chercher = ""


mainapp = tkinter.Tk()

def go_pro():
    mainapp.destroy()


def quit():
    import sys
    sys.exit()


def chercher_doc():
    file = tkinter.filedialog.askopenfile(filetypes=[("Word files", "*.docx")])
    button_go = tkinter.Button(text="Go", command=go_pro).pack()
    label_nomfichier = tkinter.Label(text=file).pack()
    global file_chercher
    file_chercher = file.name


mainapp.minsize(640, 480)
mainapp.maxsize(1280, 720)
Label_debut = tkinter.Label(mainapp, text="Choisissez le document a modifier, puis continuer").pack()
button_choix = tkinter.Button(mainapp, text="Choisir", command=chercher_doc).pack()
button_quiter = tkinter.Button(mainapp, text="Quitter", command=quit).pack()

mainapp.mainloop()


# endregion


#region Cette partie donne les informations importantes pour la suite
document_modif = Document(file_chercher)

ensemble_paragraphes = []
nbrep = 0
for par in document_modif.paragraphs:
    nvpar = fonctions.ParagraphE(par.text, nbrep)
    ensemble_paragraphes.append(nvpar)
    nbrep += 1


phrases = []
for par in ensemble_paragraphes:
    phrasest = par.phrases_original.split(".")
    nbrep = 0
    for phraseun in phrasest:
        newphrase = fonctions.PhraseE(phraseun, nbrep, par.index)
        phrases.append(newphrase)
        nbrep += 1

del nbrep
del ensemble_paragraphes


mots_trans = []
for par in phrases:
    par_t = par.phrase_final.replace(',', '')
    par_t.replace('.', '')
    mot_dans_par = par_t.split(" ")
    nBrep = 0
    for motdif in mot_dans_par:
        nvmot = fonctions.mot(motdif, par.index, nBrep)
        mots_trans.append(nvmot)
        nBrep += 1

mots = []
for mot in mots_trans:
    if len(mot.lettres_mot) > taille_min_mot:
        mots.append(mot)

del mots_trans

#endregion



#region Cette partie vas chercher les synonymes et les met dans mots_synonime[]
mots_synonime = []
fin = False

def associer_synonymes():
    optionsd = Options()
    optionsd.headless = True
    driver_synonyme = webdriver.Firefox(options=optionsd)
    for mot in mots:
        synonymes = fonctions.chercher_synonime(driver_synonyme, mot.lettres_mot)
        if synonymes != None:
            nvmotsyn = fonctions.mot_synonime(mot.lettres_mot, synonymes, mot.index_phrase, mot.index_mot)
            mots_synonime.append(nvmotsyn)
        mots.remove(mot)
    driver_synonyme.close()
    global fin
    fin = True
    print("finit")
    print(fin)


tsyno = threading.Thread(target=associer_synonymes, name="t_synonime")
tsyno.daemon = True
tsyno.start()
#endregion

#region Toute cette partie permet à l'utilisateur de de choisir les synonimes

#region Configuration de la fenêtre
root = tkinter.Tk()
root.minsize(720, 580)
root.maxsize(1100, 800)
root.geometry("720x580")

frame_principal = tkinter.Frame(root)
frame_principal.pack()

phrase_afficher = tkinter.Text(frame_principal)
phrase_afficher.pack()
phrase_afficher.tag_config('r', foreground="red")



#endregion

frame_choix_mot = tkinter.Frame(root)
frame_choix_mot.pack(fill=tkinter.BOTH)

frame_bas = tkinter.Frame(root)
frame_bas.pack(side=tkinter.BOTTOM, fill=tkinter.X)



def sauvgarder():
    document = Document()
    paragraphes_a_ajouter = []
    index_actuelle = -1

    for ph in phrases:
        if ph.index_paragraphe != index_actuelle:
            index_actuelle = ph.index_paragraphe
            paragraphes_a_ajouter.append(document.add_paragraph(ph.phrase_final))
        else:
            paragraphes_a_ajouter[ph.index_paragraphe].add_run(ph.phrase_final)

    nvdes = file_chercher.replace('.docx', '') + "-propre.docx"
    document.save(nvdes)
    import sys
    sys.exit()


boutons = []
frames_button = []


#Cette fonction permet de changer la phrase en fonction du synonime choisis
def mot_suivant(index_mot, index_synonime):
    if index_synonime != 0:
        motenquestion = mots_synonime[index_mot]
        phraseenquestion = phrases[motenquestion.index_phrase]
        phraseenquestion.changer_phrase(motenquestion.synonimes[index_synonime-1], motenquestion.index_mot)

    if fin == True and len(mots_synonime) == index_mot + 1:
        afficher_mot(index_mot)
    else:
        while len(mots_synonime) <= index_mot + 1:
            time.sleep(0.5)
        afficher_mot(index_mot + 1)

#Cette classe permet aux boutons de garder des informations
class button:
    def __init__(self, index_synon, synon, index_mot_text, frame):
        self.index_synon = index_synon
        self.synon = synon
        self.index_mot_text = index_mot_text
        self.frame = frame
        bt = tkinter.Button(self.frame, text=self.synon, command=lambda: mot_suivant(self.index_mot_text, self.index_synon))
        bt.pack(side=tkinter.LEFT)

#Cette fonction permet l'affichage de l'interface
def afficher_mot(index_mot_text):
    global boutons
    boutons = []
    global frames_button
    frames_button = []
    for widget in frame_choix_mot.winfo_children():
        widget.destroy()
    for widget in frame_bas.winfo_children():
        widget.destroy()

    motenquestion = mots_synonime[index_mot_text]
    nb_button = 1
    for synon in motenquestion.synonimes:
        if ((nb_button-1) % 8) == 0:
            frames_button.append(tkinter.Frame(frame_choix_mot))
            frames_button[-1].pack()
        if nb_button-1 == 0:
            button_ignor = tkinter.Button(frames_button[-1], text="ignorer",command=lambda: mot_suivant(index_mot_text, 0))
            button_ignor.pack(side=tkinter.LEFT)
        boutons.append(button(nb_button, synon, index_mot_text, frames_button[-1]))
        nb_button += 1


    mots_Indef = phrases[motenquestion.index_phrase].phrase_final.split(" ")
    nb_mot = 0
    phrase_afficher.config(state=tkinter.NORMAL)
    phrase_afficher.delete('1.0', tkinter.END)
    for mot in mots_Indef:
        if nb_mot < motenquestion.index_mot:
            phrase_afficher.insert(tkinter.END, mot)
            phrase_afficher.insert(tkinter.END, " ")
        elif nb_mot > motenquestion.index_mot:
            phrase_afficher.insert(tkinter.END, mot)
            phrase_afficher.insert(tkinter.END, " ")
        elif nb_mot == motenquestion.index_mot:
            phrase_afficher.insert(tkinter.INSERT, motenquestion.lettres_mot, 'r')
            phrase_afficher.insert(tkinter.END, " ")

        nb_mot += 1
    phrase_afficher.config(state=tkinter.DISABLED)
    if index_mot_text != 0:
        button_precedent = tkinter.Button(frame_bas, text="Précédent", command=lambda: afficher_mot(index_mot_text-1))
        button_precedent.pack(side=tkinter.LEFT)


    if len(mots_synonime) > index_mot_text + 1:
        button_suivant = tkinter.Button(frame_bas, text="Suivant", command=lambda: afficher_mot(index_mot_text+1))
        button_suivant.pack(side=tkinter.RIGHT)

    button_quiterroot = tkinter.Button(frame_bas, text="Sauvgarder", command=sauvgarder)
    button_quiterroot.pack()


#Cette partie permet le démarage
while len(mots_synonime)<1:
    time.sleep(0.5)

afficher_mot(0)

root.mainloop()



#endregion